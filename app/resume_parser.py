from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeParserError(Exception):
    """Erro controlado durante a leitura do currículo."""


def _normalize_text(text: str) -> str:
    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    return "\n".join(lines)


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))

        pages = []

        for page in reader.pages:
            page_text = page.extract_text() or ""

            if page_text.strip():
                pages.append(page_text)

        text = "\n".join(pages)

        if not text.strip():
            raise ResumeParserError(
                "O PDF não possui texto extraível. "
                "Ele pode ser um documento escaneado."
            )

        return _normalize_text(text)

    except ResumeParserError:
        raise

    except Exception as exc:
        raise ResumeParserError(
            "Não foi possível ler o arquivo PDF."
        ) from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))

        blocks = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                blocks.append(text)

        for table in document.tables:
            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if values:
                    blocks.append(" | ".join(values))

        text = "\n".join(blocks)

        if not text.strip():
            raise ResumeParserError(
                "O arquivo DOCX não possui texto utilizável."
            )

        return _normalize_text(text)

    except ResumeParserError:
        raise

    except Exception as exc:
        raise ResumeParserError(
            "Não foi possível ler o arquivo DOCX."
        ) from exc


def extract_resume_text(
    file_name: str,
    file_bytes: bytes,
) -> str:
    extension = Path(file_name).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParserError(
            "Formato não suportado. Envie um arquivo PDF ou DOCX."
        )

    if not file_bytes:
        raise ResumeParserError(
            "O arquivo enviado está vazio."
        )

    if extension == ".pdf":
        return _extract_pdf(file_bytes)

    if extension == ".docx":
        return _extract_docx(file_bytes)

    raise ResumeParserError(
        "Formato de currículo não reconhecido."
    )
